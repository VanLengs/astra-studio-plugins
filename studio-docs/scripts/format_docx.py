#!/usr/bin/env python3
"""
Apply Chinese government document formatting to a DOCX file.

Formats headings, body text, tables, page margins, headers, footers, and
optionally replaces mermaid placeholders with rendered images.

Requirements:
  - python-docx (`pip install python-docx`)
  - PyYAML (`pip install pyyaml`) — optional, for blueprint loading

Usage:
  python3 format_docx.py input.docx [--blueprint blueprint.yaml] [--images ./images]
"""

import argparse
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print(
        "Error: python-docx is required.\n"
        "Install it with: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


# --- Font Configuration ---

FONT_CONFIG = {
    "title": {
        "name": "方正小标宋简体",
        "name_fallback": "SimSun",
        "east_asia": "方正小标宋简体",
        "size": Pt(22),  # 二号
        "bold": True,
    },
    "heading1": {
        "name": "黑体",
        "name_fallback": "SimHei",
        "east_asia": "黑体",
        "size": Pt(16),  # 三号
        "bold": True,
    },
    "heading2": {
        "name": "楷体_GB2312",
        "name_fallback": "KaiTi",
        "east_asia": "楷体_GB2312",
        "size": Pt(16),  # 三号
        "bold": True,
    },
    "heading3": {
        "name": "仿宋_GB2312",
        "name_fallback": "FangSong",
        "east_asia": "仿宋_GB2312",
        "size": Pt(16),  # 三号
        "bold": False,
    },
    "body": {
        "name": "仿宋_GB2312",
        "name_fallback": "FangSong",
        "east_asia": "仿宋_GB2312",
        "size": Pt(12),  # 小四
        "bold": False,
    },
    "table": {
        "name": "仿宋_GB2312",
        "name_fallback": "FangSong",
        "east_asia": "仿宋_GB2312",
        "size": Pt(12),  # 小四
        "bold": False,
    },
    "table_header": {
        "name": "仿宋_GB2312",
        "name_fallback": "FangSong",
        "east_asia": "仿宋_GB2312",
        "size": Pt(12),  # 小四
        "bold": True,
    },
}

# --- Page Layout ---

PAGE_MARGINS = {
    "top": Cm(2.54),
    "bottom": Cm(2.54),
    "left": Cm(3.17),
    "right": Cm(3.17),
}

LINE_SPACING = 1.5


# --- Style Mapping ---

HEADING_STYLE_MAP = {
    "Heading 1": "heading1",
    "Heading 2": "heading2",
    "Heading 3": "heading3",
    "Title": "title",
}


def load_blueprint(blueprint_path: Path) -> dict | None:
    """Load a blueprint YAML file for additional formatting rules."""
    try:
        import yaml
    except ImportError:
        print(
            "Warning: PyYAML not installed. Blueprint loading skipped.\n"
            "Install it with: pip install pyyaml",
            file=sys.stderr,
        )
        return None

    if not blueprint_path.is_file():
        print(f"Warning: Blueprint not found: {blueprint_path}", file=sys.stderr)
        return None

    with open(blueprint_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def apply_font(run, font_config: dict):
    """Apply font settings to a run, including East Asian font mapping."""
    font = run.font
    font.size = font_config["size"]
    font.bold = font_config["bold"]

    # Set Western font
    font.name = font_config.get("name_fallback", font_config["name"])

    # Set East Asian font via XML (python-docx doesn't natively support w:eastAsia)
    r_element = run._element
    rPr = r_element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        r_element.insert(0, rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f"<w:rFonts {nsdecls('w')}/>")
        rPr.insert(0, rFonts)

    east_asia = font_config.get("east_asia", font_config["name"])
    rFonts.set(qn("w:eastAsia"), east_asia)
    rFonts.set(qn("w:ascii"), font_config.get("name_fallback", font_config["name"]))
    rFonts.set(qn("w:hAnsi"), font_config.get("name_fallback", font_config["name"]))


def set_paragraph_spacing(paragraph, line_spacing: float = LINE_SPACING):
    """Set line spacing for a paragraph."""
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f"<w:pPr {nsdecls('w')}/>")
        paragraph._element.insert(0, pPr)

    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(f"<w:spacing {nsdecls('w')}/>")
        pPr.append(spacing)

    # Line spacing: 1.5 = 360 twips (240 * 1.5)
    spacing.set(qn("w:line"), str(int(240 * line_spacing)))
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")


def format_page_layout(doc: Document):
    """Apply page margins and orientation."""
    for section in doc.sections:
        section.top_margin = PAGE_MARGINS["top"]
        section.bottom_margin = PAGE_MARGINS["bottom"]
        section.left_margin = PAGE_MARGINS["left"]
        section.right_margin = PAGE_MARGINS["right"]
        section.orientation = WD_ORIENT.PORTRAIT


def format_headings(doc: Document):
    """Apply Chinese government heading fonts to all headings."""
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name
        if style_name in HEADING_STYLE_MAP:
            config_key = HEADING_STYLE_MAP[style_name]
            config = FONT_CONFIG[config_key]

            for run in paragraph.runs:
                apply_font(run, config)

            if config_key == "title":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            set_paragraph_spacing(paragraph)


def format_body_text(doc: Document):
    """Apply body text formatting to normal paragraphs."""
    config = FONT_CONFIG["body"]

    for paragraph in doc.paragraphs:
        if paragraph.style.name in ("Normal", "Body Text", "List Paragraph"):
            for run in paragraph.runs:
                apply_font(run, config)
            set_paragraph_spacing(paragraph)


def format_tables(doc: Document):
    """Apply table formatting: borders, header bold, font size."""
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Apply full borders to the table
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = parse_xml(f"<w:tblPr {nsdecls('w')}/>")
            tbl.insert(0, tblPr)

        borders_xml = (
            f"<w:tblBorders {nsdecls('w')}>"
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            "</w:tblBorders>"
        )
        existing_borders = tblPr.find(qn("w:tblBorders"))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        tblPr.append(parse_xml(borders_xml))

        # Format cells
        for row_idx, row in enumerate(table.rows):
            is_header = row_idx == 0
            config = FONT_CONFIG["table_header"] if is_header else FONT_CONFIG["table"]

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        apply_font(run, config)


def add_page_header(doc: Document, title: str):
    """Add a page header with the document title."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        if header.paragraphs:
            p = header.paragraphs[0]
        else:
            p = header.add_paragraph()

        p.text = title
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in p.runs:
            run.font.size = Pt(9)
            run.font.name = "仿宋_GB2312"

        # Add bottom border to header
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = parse_xml(f"<w:pPr {nsdecls('w')}/>")
            p._element.insert(0, pPr)

        pBdr = parse_xml(
            f"<w:pBdr {nsdecls('w')}>"
            '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
            "</w:pBdr>"
        )
        pPr.append(pBdr)


def add_page_footer(doc: Document):
    """Add page numbers to the footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number field
        run = p.add_run()
        fld_char_begin = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
        )
        run._element.append(fld_char_begin)

        run2 = p.add_run()
        instr_text = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
        )
        run2._element.append(instr_text)

        run3 = p.add_run()
        fld_char_end = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
        )
        run3._element.append(fld_char_end)

        for r in (run, run2, run3):
            r.font.size = Pt(9)


def replace_mermaid_with_images(doc: Document, images_dir: Path):
    """Replace mermaid code block placeholders with rendered images."""
    if not images_dir.is_dir():
        return

    mermaid_pattern = re.compile(r"```mermaid.*?```", re.DOTALL)
    figure_pattern = re.compile(r"!\[(.+?)\]\((.+?)\)")

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text

        # Look for image references that point to our images directory
        fig_match = figure_pattern.search(text)
        if fig_match:
            caption = fig_match.group(1)
            img_ref = fig_match.group(2)

            # Resolve image path
            img_path = images_dir / os.path.basename(img_ref)
            if not img_path.is_file():
                img_path = Path(img_ref)

            if img_path.is_file():
                # Clear the paragraph and insert the image
                paragraph.clear()
                run = paragraph.add_run()
                try:
                    run.add_picture(str(img_path), width=Cm(14.5))
                except Exception as e:
                    run.add_text(f"[Image: {caption} — rendering failed: {e}]")

                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Add caption as a new paragraph after (if not already present)
                # Note: modifying doc.paragraphs during iteration is tricky,
                # so we just append caption text to the same paragraph
                caption_run = paragraph.add_run(f"\n{caption}")
                caption_run.font.size = Pt(9)
                caption_run.font.italic = True


def extract_title(doc: Document) -> str:
    """Extract the document title from the first heading or title paragraph."""
    for paragraph in doc.paragraphs:
        if paragraph.style.name in ("Title", "Heading 1") and paragraph.text.strip():
            return paragraph.text.strip()
    # Fallback: first non-empty paragraph
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            return paragraph.text.strip()
    return "项目文档"


def format_document(
    input_path: Path,
    output_path: Path | None = None,
    blueprint_path: Path | None = None,
    images_dir: Path | None = None,
):
    """Main formatting pipeline."""
    if not input_path.is_file():
        print(f"Error: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    print(f"Loading: {input_path}")
    doc = Document(str(input_path))

    # Load blueprint if provided
    blueprint = None
    if blueprint_path:
        blueprint = load_blueprint(blueprint_path)
        if blueprint:
            print(f"Blueprint: {blueprint_path}")

    # Extract title for page header
    title = extract_title(doc)
    print(f"Title: {title}")

    # Apply formatting pipeline
    print("Applying page layout...")
    format_page_layout(doc)

    print("Formatting headings...")
    format_headings(doc)

    print("Formatting body text...")
    format_body_text(doc)

    print("Formatting tables...")
    format_tables(doc)

    print("Adding page header...")
    add_page_header(doc, title)

    print("Adding page footer (page numbers)...")
    add_page_footer(doc)

    # Replace mermaid images if directory provided
    if images_dir:
        print(f"Replacing mermaid diagrams from: {images_dir}")
        replace_mermaid_with_images(doc, images_dir)

    # Save
    out = output_path or input_path
    doc.save(str(out))
    print(f"\nSaved: {out}")


def main():
    parser = argparse.ArgumentParser(
        description="Apply Chinese government document formatting to a DOCX file.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  python3 format_docx.py input.docx\n"
            "  python3 format_docx.py input.docx --blueprint blueprint.yaml\n"
            "  python3 format_docx.py input.docx --images ./images\n"
            "  python3 format_docx.py input.docx --blueprint bp.yaml --images ./images\n"
        ),
    )
    parser.add_argument("input", type=Path, help="Input DOCX file path")
    parser.add_argument(
        "--output", "-o",
        type=Path,
        default=None,
        help="Output DOCX file path (default: overwrite input)",
    )
    parser.add_argument(
        "--blueprint",
        type=Path,
        default=None,
        help="Blueprint YAML file for additional formatting rules",
    )
    parser.add_argument(
        "--images",
        type=Path,
        default=None,
        help="Directory containing rendered Mermaid images",
    )

    args = parser.parse_args()
    format_document(args.input, args.output, args.blueprint, args.images)


if __name__ == "__main__":
    main()
